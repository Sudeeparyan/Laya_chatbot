"""Generate the synthetic evaluation corpus and ingest it through the real pipeline.

    python backend/scripts/build_mock_corpus.py [--keep-existing]

The documents are written as Word files and then handed to ``convert_upload`` —
the same function the ``POST /api/convert`` endpoint calls. Nothing here writes
Markdown or a manifest by hand, so the corpus the retriever indexes is the
output of the application's own converter rather than a fixture that merely
looks like one. If the converter's behaviour changes, this corpus changes with
it, which is the only way the evaluation numbers stay honest.

Determinism matters because ``document_id`` is derived from the file's SHA-256:
a corpus rebuilt from identical source text must produce identical ids, or every
graph snapshot, saved question set and screenshot in the write-up goes stale for
no reason. Word files are ZIP archives and ``python-docx`` stamps each entry
with the current time, so the archive is rewritten with a fixed timestamp before
it is hashed (see ``_normalise_archive``).
"""
from __future__ import annotations

import argparse
import shutil
import sys
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from tempfile import TemporaryDirectory

BACKEND_DIR = Path(__file__).resolve().parents[1]
if str(BACKEND_DIR) not in sys.path:
    sys.path.insert(0, str(BACKEND_DIR))

from docx import Document  # noqa: E402
from docx.shared import Pt  # noqa: E402

from app.config import (  # noqa: E402
    MANIFEST_DIR,
    MARKDOWN_OUTPUT_DIR,
    RAW_UPLOAD_DIR,
    ensure_data_dirs,
)
from app.services.markdown_converter import convert_upload  # noqa: E402

from mock_corpus_content import ALL_DOCUMENTS, MockDocument  # noqa: E402

#: Fixed archive timestamp, so a rebuild from identical text hashes identically.
FIXED_ZIP_TIME = (2026, 1, 1, 0, 0, 0)

#: Written into every document so a reader who opens one outside the app still
#: sees what it is. It survives conversion into the Markdown the agent reads.
PROVENANCE = (
    "Synthetic document. Scheme names, benefit values, limits and conditions are "
    "invented for research use and are not any organisation's approved wording."
)


def _render_docx(document: MockDocument, target: Path) -> None:
    """Write one mock document as a Word file."""
    docx = Document()

    style = docx.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # The provenance notice goes in the page footer and the file properties,
    # never in the body.
    #
    # Anyone who opens the Word file sees it on every page, which is the point.
    # But it is byte-identical across all nine documents, and body text that
    # repeats verbatim in every document is poison for this pipeline: the term
    # miner scores it as a genuine multi-document concept ("synthetic document",
    # "research use"), and the TF-IDF pass reads the shared paragraph as
    # evidence that every document resembles every other. Footers are not
    # extracted by the converter, so the notice stays visible to a human and
    # invisible to the index.
    footer = docx.sections[0].footer.paragraphs[0]
    footer.text = PROVENANCE
    docx.core_properties.comments = PROVENANCE
    docx.core_properties.category = "Synthetic research corpus"

    docx.add_heading(document.title, level=1)
    docx.add_paragraph(document.intro)

    for section in document.sections:
        docx.add_heading(section.heading, level=2)
        for paragraph in section.paragraphs:
            docx.add_paragraph(paragraph)
        if section.table:
            header, *rows = section.table
            table = docx.add_table(rows=1, cols=len(header))
            table.style = "Table Grid"
            for cell, text in zip(table.rows[0].cells, header):
                cell.text = text
            for row in rows:
                cells = table.add_row().cells
                for cell, text in zip(cells, row):
                    cell.text = text

    # python-docx sets these to "now" on save; pinning them keeps the bytes stable.
    epoch = datetime(2026, 1, 1, tzinfo=timezone.utc).replace(tzinfo=None)
    docx.core_properties.created = epoch
    docx.core_properties.modified = epoch
    docx.core_properties.author = document.owner
    docx.core_properties.last_modified_by = document.owner
    docx.core_properties.revision = 1

    docx.save(target)
    _normalise_archive(target)


def _normalise_archive(path: Path) -> None:
    """Rewrite a .docx with fixed entry timestamps so its hash is reproducible."""
    with zipfile.ZipFile(path) as archive:
        entries = [(info.filename, archive.read(info.filename)) for info in archive.infolist()]

    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as archive:
        for name, payload in entries:
            info = zipfile.ZipInfo(name, date_time=FIXED_ZIP_TIME)
            info.compress_type = zipfile.ZIP_DEFLATED
            info.external_attr = 0o600 << 16
            archive.writestr(info, payload)


def _clear_corpus() -> list[str]:
    """Empty the three corpus directories, reporting what went.

    The previous corpus is removed rather than added to. Leaving a stale
    document in place would leave it in the index and in every count the
    write-up quotes, which is the kind of drift that is very hard to notice
    later.
    """
    removed: list[str] = []
    for directory, patterns in (
        (RAW_UPLOAD_DIR, ("*",)),
        (MARKDOWN_OUTPUT_DIR, ("*.md",)),
        (MANIFEST_DIR, ("*.json",)),
    ):
        if not directory.exists():
            continue
        for pattern in patterns:
            for path in sorted(directory.glob(pattern)):
                if path.is_file():
                    path.unlink()
                    removed.append(path.name)
    return removed


def build(keep_existing: bool = False) -> list[dict[str, object]]:
    ensure_data_dirs()

    if not keep_existing:
        removed = _clear_corpus()
        if removed:
            print(f"Cleared {len(removed)} existing corpus file(s).")

    written: list[dict[str, object]] = []

    with TemporaryDirectory(prefix="mock-corpus-") as staging_name:
        staging = Path(staging_name)

        for document in ALL_DOCUMENTS:
            source = staging / document.filename
            _render_docx(document, source)

            result = convert_upload(
                source,
                original_filename=document.filename,
                document_title=document.title,
                department=document.department,
                access_group=document.access_group,
                classification=document.classification,
                document_owner=document.owner,
                version=document.version,
                # Deliberately off: a vision call would make the corpus depend
                # on a model endpoint and stop being reproducible offline.
                enable_plugin=False,
            )

            markdown_path = Path(result.metadata.markdown_path)
            written.append(
                {
                    "document_id": result.metadata.document_id,
                    "title": document.title,
                    "department": document.department,
                    "chars": len(result.markdown),
                    "confidence": result.confidence.overall,
                    "markdown": markdown_path.name,
                }
            )
            print(f"  {result.metadata.document_id:<52} {len(result.markdown):>6} chars")

    return written


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--keep-existing",
        action="store_true",
        help="Add to the current corpus instead of replacing it.",
    )
    args = parser.parse_args()

    print(f"Building {len(ALL_DOCUMENTS)} synthetic documents through the real converter…")
    written = build(keep_existing=args.keep_existing)

    total_chars = sum(int(item["chars"]) for item in written)
    print(
        f"\nWrote {len(written)} document(s), {total_chars:,} characters of Markdown."
        f"\n  raw uploads     {RAW_UPLOAD_DIR}"
        f"\n  markdown        {MARKDOWN_OUTPUT_DIR}"
        f"\n  manifests       {MANIFEST_DIR}"
        "\n\nNext: python backend/scripts/build_knowledge_graph.py"
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
